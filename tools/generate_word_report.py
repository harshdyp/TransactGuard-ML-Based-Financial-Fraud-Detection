from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _iter_markdown_blocks(md: str):
    # Very small markdown-to-docx helper:
    # - Headings: #, ##, ### -> doc heading levels 1..3
    # - Bullets: lines starting with "- " -> bullet paragraphs
    # - Everything else -> normal paragraphs (blank lines separate paragraphs)
    lines = md.splitlines()
    para_buf: list[str] = []

    def flush_para():
        nonlocal para_buf
        if para_buf:
            yield ("p", "\n".join(para_buf).strip())
            para_buf = []

    for line in lines:
        raw = line.rstrip()
        if not raw.strip():
            yield from flush_para()
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", raw)
        if m:
            yield from flush_para()
            level = min(len(m.group(1)), 6)
            yield ("h", level, m.group(2).strip())
            continue

        if raw.startswith("- "):
            yield from flush_para()
            yield ("li", raw[2:].strip())
            continue

        para_buf.append(raw)

    yield from flush_para()


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    md = md_path.read_text(encoding="utf-8")
    doc = Document()

    for block in _iter_markdown_blocks(md):
        if not block:
            continue

        if block[0] == "h":
            _, level, text = block
            # Map markdown headings to Word headings (Title is handled separately if needed)
            heading_level = max(1, min(int(level), 4))
            doc.add_heading(text, level=heading_level)
            continue

        if block[0] == "li":
            _, text = block
            doc.add_paragraph(text, style="List Bullet")
            continue

        if block[0] == "p":
            _, text = block
            p = doc.add_paragraph(text)
            # Center the first title line if it looks like a cover heading
            if text.strip().lower().startswith("final project report"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

    doc.save(str(docx_path))


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    md_path = repo_root / "Final_Project_Report_FraudX.md"
    out_path = repo_root / "Final_Project_Report_FraudX.docx"

    if not md_path.exists():
        raise SystemExit(f"Missing input file: {md_path}")

    md_to_docx(md_path, out_path)
    print(f"Created: {out_path}")


if __name__ == "__main__":
    main()

